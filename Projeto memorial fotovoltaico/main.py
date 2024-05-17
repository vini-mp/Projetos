#Bibliotecas
#pip install docx2pdf
#pip install openpyxl
#pip install python-docx

from docx import Document
from docx.shared import Pt 
from docx.shared import RGBColor 
from docx2pdf import convert
from openpyxl import load_workbook
import datetime
from data_conv import data
import warnings
warnings.filterwarnings("ignore")

#Abrindo arquivo Word
arquivoWord=Document("Memorial.docx")

 #Configurando os estilos do documento
estilo=arquivoWord.styles["Normal"]


#Configurando o caminho do arquivo excel
nome_arquivos_alunos="DadosCliente.xlsm"

#Abrindo o arquivo Excel
planilhaDadosAlunos=load_workbook(nome_arquivos_alunos)

#Selecionando a aba da planilha que iremos trabalhar
sheet_selecionada=planilhaDadosAlunos["Cliente"]


#Configutando estilo geral
fonte=estilo.font
fonte.name="Arial"
fonte.size=Pt(18)


#Obtendo dados do arquivo Excel
nomeCliente=sheet_selecionada['B2'].value
cpf=sheet_selecionada['B3'].value
numeroInstalacao=sheet_selecionada['B4'].value
municipio=sheet_selecionada['B5'].value
potenciaModulo = sheet_selecionada['B12'].value
potenciaInversor = sheet_selecionada['B13'].value
quantidadePlacas = sheet_selecionada['B18'].value
quantidadeInversores = sheet_selecionada['B19'].value
potenciaUnitariaPlaca = sheet_selecionada['B20'].value
potenciaUnitariaInversor = sheet_selecionada['B21'].value
micro=sheet_selecionada['B23'].value
rua=sheet_selecionada['B6'].value
numero=sheet_selecionada['B7'].value
bairro=sheet_selecionada['B8'].value
complemento=sheet_selecionada['B9'].value
cep=sheet_selecionada['B10'].value
modeloInversor=sheet_selecionada['B17'].value
modeloPlaca=sheet_selecionada['B16'].value
fabricanteModulo=sheet_selecionada['B14'].value
fabricanteInversor=sheet_selecionada['B15'].value
arranjo=sheet_selecionada['B22'].value
nomeEng=sheet_selecionada['B29'].value
cpfEng=sheet_selecionada['B30'].value
concessionaria=sheet_selecionada['B33'].value
ncrea = sheet_selecionada['B31'].value
estrutura=sheet_selecionada['B26'].value
cabeamentocc=sheet_selecionada['B24'].value
cabeamentoca=sheet_selecionada['B25'].value
stringbox=sheet_selecionada['B27'].value


#Controle se o inversor utilizado é da classe dos microinversores
if micro=="SIM":
    micro_controle=" micro "     

else:
     micro_controle=" "
     

#Manipulação do arquivo Word para a inclusão das informações do cliente
for paragrafo in arquivoWord.paragraphs:

        if "@conce" in paragrafo.text:
                                        
            frase_montada = (f"O presente memorial tem por finalidade indicar os materiais e serviços a serem  "
                             f"aplicados na instalação de sistema fotovoltaico, seguindo os critérios das "
                             f"resoluções ANEEL 482/2011 e 687/2015, Norma de Fornecimento da {concessionaria} e "
                             f"Especificações Técnicas de Materiais e Serviços.")

                    
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)


        if "@nome" in paragrafo.text:
            
            frase_montada = f"Nome do Cliente: {nomeCliente}"
            
            paragrafo.text= ""

            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

            adicionaNovaPalavra.bold=True
            


        if "@cpf" in paragrafo.text:
            
            frase_montada = f"CPF: {cpf}"
            
            paragrafo.text= ""
            
            
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

            
            adicionaNovaPalavra.bold=True
            


        if "@instalacao" in paragrafo.text:
            
            frase_montada = f"Número da instalação: {numeroInstalacao}"
            
            paragrafo.text= ""

            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

            adicionaNovaPalavra.bold=True
            

        if "@municipio" in paragrafo.text:
            
            frase_montada = f"Município: {municipio}"
            
            paragrafo.text= ""

            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

            adicionaNovaPalavra.bold=True

        
        if "@capacidade" in paragrafo.text:
                                 
            if arranjo==1:

                controle = input("Digite o nível de proteção do disjuntor: " )

                                
                frase_montada = (f"Geração de {potenciaModulo} kW de potência de pico e {potenciaInversor}kW de "
                    f"potência ativa geradas através de {quantidadePlacas} módulo(s) de {potenciaUnitariaPlaca}W ligado(s) "
                    f"em {quantidadeInversores} {micro_controle} inversor(s) de {potenciaInversor}kW. Foi instalada proteção CA "
                    f"com um disjuntor de {controle}A.")

                
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada)

            if arranjo>1:
                controle = input("Digite o nível de proteção dos disjuntores (separados por vírgula): " )
                               
                
                frase_montada = (f"Geração de {potenciaModulo} kW de potência de pico e {potenciaInversor}kW de "
                    f"potência ativa geradas através de {quantidadePlacas} módulo(s) de {potenciaUnitariaPlaca}W ligado(s) "
                    f"em {quantidadeInversores} {micro_controle} inversor(s) de {potenciaInversor}kW. Foram instaladas proteções CA "
                    f"com disjuntores de {controle}A para cada arranjo.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada)
                adicionaNovaPalavra.font.size=Pt(12)


        
        if "@endereco" in paragrafo.text:
                                 
            frase_montada = (f"A instalação fotovoltaica será realizada na instalação situada no "
                f"endereço {rua}, {numero}, {complemento}, bairro: "
                f"{bairro}, município de {municipio}, no estado de Minas Gerais,  "
                f"CEP: {cep}.")

            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

        

        if "@modulo" in paragrafo.text:
            area=(quantidadePlacas*2.5) +10                   
            frase_montada = (f"Fabricante: {fabricanteModulo} \n Modelo: {modeloPlaca} \n Quantidade de módulos: {quantidadePlacas} \n Área dos arranjos (m²): {area} \n Potência máxima por módulo: {potenciaUnitariaPlaca}W")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)



        if "@inversor" in paragrafo.text:
                              
            frase_montada = (f"Fabricante: {fabricanteInversor} \n Modelo: {modeloInversor} \n Quantidade de inversores: {quantidadeInversores} \n Potência máxima de saída: {potenciaUnitariaInversor}kW \n Fator de potência: >0,99")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)

        
        if "@obra" in paragrafo.text:
            area=(quantidadePlacas*2.5) +10                    
            frase_montada = (f"Área mínima que o sistema ocupará é de {area} m². ")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)


        if "@arranjo" in paragrafo.text:
            
            if arranjo==1:
                frase_montada1 = (f"Arranjo 1: {quantidadePlacas} módulos {modeloPlaca} ligados ao {micro_controle} inversor {modeloInversor}.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)


            if arranjo>1:
                
                controle = input("Digite a quantidade de placas por arranjo (separadas por vírgula): " )
                controle2 = input("Digite a quantidade de inversores por arranjo (separadas por vírgula): " )
                lista=[]
                i=0
                while i<arranjo:
                    lista.append(i+1)
                    i+=1
                
                
                
                frase_montada1 = (f"Arranjos {lista} com {controle} módulos {modeloPlaca} e {controle2} {micro_controle}inverores {modeloInversor}, respectivamente.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)
            

        if "@equipamento" in paragrafo.text:
            
            frase_montada = (f"Serão utilizados {quantidadeInversores}{micro_controle}inversor(s) da marca {fabricanteInversor} operando em 220 V (CA) com  "
            f"potência de {potenciaInversor} kW. Não haverá uso de autotransformador já que a instalação "
            f"possui a tensão nominal requerida.")
                
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)


                        
        if "@data" in paragrafo.text:
            
                data_hora_atual = datetime.datetime.now()

                #  Extrair o dia, o mês e o ano
                dia = data_hora_atual.day
                mes = data_hora_atual.month
                ano = data_hora_atual.year

                #Função para verificar qual é mês em português 
                mes_nome=data(mes)               
                
                frase_montada1 = (f"{municipio}, {dia} de {mes_nome} de {ano}")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)

        
        if "@engenheiro" in paragrafo.text:
            
            frase_montada1 = (f"Eng.° {nomeEng} ")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada1)

            
        if "@cpfengenheiro" in paragrafo.text:
            
            frase_montada1 = (f"CPF: {cpfEng} ")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
            adicionaNovaPalavra.bold=False
            

        if "@crea" in paragrafo.text:
                                        
            frase_montada1 = (f"CREA: {ncrea} ")
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
            adicionaNovaPalavra.bold=False


        if "@cabos" in paragrafo.text:
            
            frase_montada = (f"Serão utilizados cabos solares com proteção UV de {cabeamentocc} mm². As conexões serão "
            f"feitas por conectores MC4 com proteção UV e resistência a amoníaco. "
            f"Serão utilizados cabos de {cabeamentoca} mm² para o lado CA.")
                
                            
            paragrafo.text= ""
            adicionaNovaPalavra = paragrafo.add_run(frase_montada)


        if "@string" in paragrafo.text:
            
            if stringbox=="NÃO":
                frase_montada1 = (f"Não haverá String Box externa. O DPS e a chave seccionadora são integradas ao inversor.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)


            if stringbox=="SIM":
                
                controle = input("Digite a quantidade de DPS utilizados: " )
                controle2 = input("Digite a quantidade de fusíveis/disjuntores utilizados): " )
                controle3 = input("Digite a capacidade de proteção dos fusíveis/disjuntores utilizados): " )
                                
                frase_montada1 = (f"Serão empregados {controle} DPS e {controle2} fusíveis/disjuntores com capacidade de "
                                  f"proteção de {controle3}A para comporem a string box CC.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)

        
        if "@estrutura" in paragrafo.text:
            
            if estrutura=="TELHADO":
                
                frase_montada1 = (f"Serão utilizados trilhos em alumínio para fixação dos painéis nas telhas da edificação.")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)


            if estrutura=="SOLO":
                       
                controle = input("Digite a forma como os painéis serão fixos ao solo: Serão empregados... " )
                frase_montada1 = (f"Os painéis serão instalados em solo.Serão empregados {controle}")
                            
                paragrafo.text= ""
                adicionaNovaPalavra = paragrafo.add_run(frase_montada1)
                adicionaNovaPalavra.font.size=Pt(12)

    
#Caminho do arquivo a ser salvo
caminhoArquivo=nomeCliente+".docx"

#Salva o arquivo com o nome do cliente
arquivoWord.save(caminhoArquivo)

convert(caminhoArquivo)

print("Arquivo gerado com sucesso!")

